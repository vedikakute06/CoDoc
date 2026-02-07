"""
Document generation utilities for the Code Documentation Generator.

This module focuses on turning analysis results into human-friendly documents:
- LeetCode-style solution reports in DOCX format
- GitHub repository READMEs in Markdown
"""

from __future__ import annotations

import datetime as _dt
from pathlib import Path
from typing import Any, Dict

from docx import Document  # type: ignore[import]
from docx.enum.style import WD_STYLE_TYPE  # type: ignore[import]
from docx.shared import Pt  # type: ignore[import]


class DocGenerator:
    """
    Helper class for generating documentation artifacts from analysis data.
    """

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------
    @staticmethod
    def _ensure_parent_dir(output_path: Path) -> None:
        output_path.parent.mkdir(parents=True, exist_ok=True)

    @staticmethod
    def _get_or_create_code_style(document: Document) -> str:
        """
        Ensure there is a monospace style available for code snippets.
        Returns the style name to use.
        """
        styles = document.styles

        # Try to reuse an existing style
        if "Code" in styles:
            return "Code"

        # Otherwise create a simple monospace paragraph style
        code_style = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
        font = code_style.font
        font.name = "Courier New"
        font.size = Pt(9)
        return "Code"

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------
    def generate_leetcode_report(self, analysis_data: Dict[str, Any], output_path: str) -> Path:
        """
        Generate a professional-looking DOCX report for a LeetCode-style solution.

        Expects `analysis_data` with keys:
        - original_code (str)
        - language (str)
        - analysis (dict)            -> from LLMHandler.analyze_code
        - optimized_version (dict)   -> from LLMHandler.get_optimized_version
        - alternatives (dict)        -> from LLMHandler.get_alternative_approaches
        """
        path = Path(output_path)
        if path.suffix.lower() != ".docx":
            path = path.with_suffix(".docx")
        self._ensure_parent_dir(path)

        doc = Document()
        code_style_name = self._get_or_create_code_style(doc)

        # Title
        title = doc.add_heading(level=0)
        run = title.add_run("LeetCode Solution Analysis Report")
        run.bold = True

        # Date
        date_par = doc.add_paragraph()
        date_par.add_run(f"Generated on: {_dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        # Metadata
        language = analysis_data.get("language", "Unknown")
        doc.add_paragraph(f"Language: {language}")

        # Original Code
        doc.add_heading("Original Code", level=1)
        code_par = doc.add_paragraph(style=code_style_name)
        code_par.add_run(analysis_data.get("original_code", "").rstrip() + "\n")

        # Analysis
        analysis = analysis_data.get("analysis", {}) or {}
        doc.add_heading("Analysis", level=1)

        desc = analysis.get("description") or "N/A"
        tc = analysis.get("time_complexity") or "N/A"
        sc = analysis.get("space_complexity") or "N/A"
        score = analysis.get("code_quality_score")
        issues = analysis.get("issues") or []
        suggestions = analysis.get("improvement_suggestions") or []

        doc.add_heading("What the Code Does", level=2)
        doc.add_paragraph(str(desc))

        doc.add_heading("Complexity", level=2)
        doc.add_paragraph(f"Time Complexity: {tc}")
        doc.add_paragraph(f"Space Complexity: {sc}")

        doc.add_heading("Code Quality", level=2)
        doc.add_paragraph(f"Code Quality Score (1–10): {score if score is not None else 'N/A'}")

        doc.add_heading("Issues / Potential Bugs", level=2)
        if issues:
            for item in issues:
                doc.add_paragraph(str(item), style="List Bullet")
        else:
            doc.add_paragraph("No major issues identified or not provided.")

        doc.add_heading("Improvement Suggestions", level=2)
        if suggestions:
            for item in suggestions:
                doc.add_paragraph(str(item), style="List Bullet")
        else:
            doc.add_paragraph("No specific suggestions provided.")

        # Optimized Version
        optimized = analysis_data.get("optimized_version", {}) or {}
        doc.add_heading("Optimized Version", level=1)
        opt_code = optimized.get("optimized_code") or optimized.get("raw") or ""
        if opt_code:
            opt_par = doc.add_paragraph(style=code_style_name)
            opt_par.add_run(str(opt_code).rstrip() + "\n")
        else:
            doc.add_paragraph("No optimized version provided.")

        opt_explanation = optimized.get("explanation")
        if opt_explanation:
            doc.add_heading("Optimization Explanation", level=2)
            doc.add_paragraph(str(opt_explanation))

        # Alternative Approaches
        alternatives = analysis_data.get("alternatives", {}) or {}
        doc.add_heading("Alternative Approaches", level=1)
        approaches = alternatives.get("approaches") or []
        if approaches:
            for idx, approach in enumerate(approaches, start=1):
                name = approach.get("name") or f"Approach {idx}"
                doc.add_heading(name, level=2)
                doc.add_paragraph(f"Description: {approach.get('description') or 'N/A'}")
                doc.add_paragraph(
                    f"Time Complexity: {approach.get('time_complexity') or 'N/A'}"
                )
                doc.add_paragraph(
                    f"When to Use: {approach.get('when_to_use') or 'N/A'}"
                )
        else:
            doc.add_paragraph("No alternative approaches provided.")

        doc.save(path)
        return path

    def generate_readme(
        self,
        analysis_data: Dict[str, Any],
        repo_info: Dict[str, Any],
        output_path: str,
    ) -> Path:
        """
        Generate a README.md file based on repository info and analysis data.

        `analysis_data` is expected to have:
        - overview
        - tech_stack
        - installation
        - usage
        - contributing

        `repo_info` should contain:
        - name
        - description
        - stars
        - forks
        - language (optional)
        - topics (optional)
        """
        path = Path(output_path)
        if path.suffix.lower() != ".md":
            path = path.with_suffix(".md")
        self._ensure_parent_dir(path)

        name = repo_info.get("name", "Project")
        description = repo_info.get("description", "")
        stars = repo_info.get("stars", 0)
        forks = repo_info.get("forks", 0)
        language = repo_info.get("language") or "Unknown"
        topics = repo_info.get("topics") or []

        overview = analysis_data.get("overview", "").strip()
        tech_stack = analysis_data.get("tech_stack", "").strip()
        installation = analysis_data.get("installation", "").strip()
        usage = analysis_data.get("usage", "").strip()
        contributing = analysis_data.get("contributing", "").strip()

        # Simple shields-style badges using static counts
        badges = []
        badges.append(f"![Stars](https://img.shields.io/badge/stars-{stars}-brightgreen)")
        badges.append(f"![Forks](https://img.shields.io/badge/forks-{forks}-blue)")
        badges.append(f"![Language](https://img.shields.io/badge/language-{language}-orange)")

        lines = []
        lines.append(f"# {name}")
        if badges:
            lines.append("")
            lines.append(" ".join(badges))
        if description:
            lines.append("")
            lines.append(description)

        if overview:
            lines.append("")
            lines.append("## Overview")
            lines.append("")
            lines.append(overview)

        if topics:
            lines.append("")
            lines.append("## Tags")
            lines.append("")
            lines.append(", ".join(str(t) for t in topics))

        if tech_stack:
            lines.append("")
            lines.append("## Tech Stack")
            lines.append("")
            lines.append(tech_stack)

        if installation:
            lines.append("")
            lines.append("## Installation")
            lines.append("")
            lines.append(installation)

        if usage:
            lines.append("")
            lines.append("## Usage")
            lines.append("")
            lines.append(usage)

        if contributing:
            lines.append("")
            lines.append("## Contributing")
            lines.append("")
            lines.append(contributing)

        # Simple default license section
        lines.append("")
        lines.append("## License")
        lines.append("")
        lines.append("This project is licensed under the terms specified in the repository.")

        content = "\n".join(lines).rstrip() + "\n"
        path.write_text(content, encoding="utf-8")
        return path


if __name__ == "__main__":
    """
    Simple manual test for DocGenerator.

    Run from project root:
        python -m backend.doc_generator
    """

    # Sample LeetCode-style analysis data
    sample_leetcode_analysis = {
        "original_code": """def two_sum(nums, target):
    for i in range(len(nums)):
        for j in range(i + 1, len(nums)):
            if nums[i] + nums[j] == target:
                return [i, j]
    return []""",
        "language": "Python",
        "analysis": {
            "description": "Finds indices of two numbers in the list that add up to the target.",
            "time_complexity": "O(n^2)",
            "space_complexity": "O(1)",
            "code_quality_score": 6,
            "issues": ["Inefficient for large lists due to nested loops."],
            "improvement_suggestions": ["Use a hash map to achieve O(n) time complexity."],
        },
        "optimized_version": {
            "optimized_code": """def two_sum(nums, target):
    index = {}
    for i, num in enumerate(nums):
        diff = target - num
        if diff in index:
            return [index[diff], i]
        index[num] = i
    return []""",
            "explanation": "Uses a hash map to store seen values, reducing time complexity to O(n).",
        },
        "alternatives": {
            "approaches": [
                {
                    "name": "Two-pointer (sorted)",
                    "description": "Sort the array and use two pointers from both ends.",
                    "time_complexity": "O(n log n) due to sorting",
                    "when_to_use": "When modifying the order of elements is acceptable.",
                }
            ]
        },
    }

    # Sample GitHub repo analysis data
    sample_repo_info = {
        "name": "sample-repo",
        "description": "A small sample repository used for testing the DocGenerator.",
        "stars": 42,
        "forks": 7,
        "language": "Python",
        "topics": ["documentation", "example", "testing"],
    }
    sample_repo_analysis = {
        "overview": "This project demonstrates how to generate documentation from analysis data.",
        "tech_stack": "- Python 3.10+\n- FastAPI\n- PostgreSQL",
        "installation": "Clone the repository and run `pip install -r requirements.txt`.",
        "usage": "Run `python main.py` and open the provided URL in your browser.",
        "contributing": "Fork the repo, create a feature branch, and open a pull request.",
    }

    generator = DocGenerator()

    # Generate LeetCode report DOCX
    leetcode_report_path = generator.generate_leetcode_report(
        sample_leetcode_analysis, "output/sample_leetcode_report.docx"
    )
    print(f"Generated LeetCode report at: {leetcode_report_path}")

    # Generate README.md
    readme_path = generator.generate_readme(
        sample_repo_analysis, sample_repo_info, "output/sample_README.md"
    )
    print(f"Generated README at: {readme_path}")

